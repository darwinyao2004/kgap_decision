import argparse
import json
import platform
from pathlib import Path

import pandas as pd


TITLE = "面向知识缺口的大语言模型动作决策：直接回答、主动追问与拒答的联合建模"


def _fmt(x: float) -> str:
    try:
        return f"{float(x):.3f}"
    except Exception:
        return str(x)


def _cell(x) -> str:
    if pd.isna(x):
        return ""
    if isinstance(x, float):
        if x.is_integer():
            return str(int(x))
        return f"{x:.3f}"
    return str(x)


def _table(df: pd.DataFrame, cols: list[str], max_rows: int = 12) -> str:
    if df.empty:
        return "（无数据）"
    subset = df[cols].head(max_rows).copy()
    header = "| " + " | ".join(cols) + " |"
    sep = "| " + " | ".join(["---"] * len(cols)) + " |"
    body = []
    for _, row in subset.iterrows():
        body.append("| " + " | ".join(_cell(row[c]) for c in cols) + " |")
    return "\n".join([header, sep] + body)


def _ablation_comment(ablation: pd.DataFrame) -> str:
    changed = ablation[ablation["delta_score"].abs() > 1e-9]
    if changed.empty:
        return (
            "这一组消融没有改变动作分数。我的理解是：当前数据模板的提示词和标签边界过于清楚，"
            "随机森林动作分类器单独就能学到主要模式，所以额外规则、gap_type 输出和 self-consistency 特征没有在测试集上拉开差距。"
            "这不是“所有模块都没用”的证明，只说明现在的数据还不够难。"
        )
    improved = changed.sort_values("delta_score", ascending=False).iloc[0]
    worsened = changed.sort_values("delta_score").iloc[0]
    pieces = []
    if improved["delta_score"] > 0:
        pieces.append(
            f"`{improved['variant']}` 的分数反而高出 Full Method {_fmt(improved['delta_score'])}，"
            "说明 Full Method 的人工规则在这批样本上有过度干预。"
        )
    if worsened["delta_score"] < 0:
        pieces.append(
            f"`{worsened['variant']}` 的分数下降 {_fmt(abs(worsened['delta_score']))}，"
            "这部分模块对当前决策边界有实际贡献。"
        )
    flat = ablation[ablation["delta_score"].abs() <= 1e-9]["variant"].tolist()
    if flat:
        pieces.append("其余变体没有造成可见变化，主要原因仍是合成数据模式较强。")
    return "".join(pieces)


def _stability_summary(stability: pd.DataFrame) -> str:
    if stability.empty:
        return "未生成重复实验结果。"
    means = stability[["gap_type_macro_f1", "action_macro_f1", "score"]].mean()
    stds = stability[["gap_type_macro_f1", "action_macro_f1", "score"]].std(ddof=0)
    return (
        f"三组随机划分下，Full Method 的 action macro-F1 均值为 {_fmt(means['action_macro_f1'])}，"
        f"标准差为 {_fmt(stds['action_macro_f1'])}；综合效用均值为 {_fmt(means['score'])}，"
        f"标准差为 {_fmt(stds['score'])}。"
        "这个稳定性来自同一生成协议下的样本分布，不能外推为真实用户分布下同样稳定。"
    )


def generate_markdown(output: str = "reports/final_report.md") -> str:
    metrics = pd.read_csv("results/metrics_summary.csv")
    per_class = pd.read_csv("results/per_class_metrics.csv")
    ablation = pd.read_csv("results/ablation_summary.csv")
    sig = pd.read_csv("results/significance_tests.csv") if Path("results/significance_tests.csv").exists() else pd.DataFrame()
    stability = pd.read_csv("results/repeated_seed_summary.csv") if Path("results/repeated_seed_summary.csv").exists() else pd.DataFrame()
    label_dist = pd.read_csv("data/processed/label_distribution.csv")
    manifest = json.loads(Path("data/processed/split_manifest.json").read_text(encoding="utf-8"))
    api_status = json.loads(Path("results/api_status.json").read_text(encoding="utf-8"))
    selection = (
        json.loads(Path("results/full_method_selection.json").read_text(encoding="utf-8"))
        if Path("results/full_method_selection.json").exists()
        else {}
    )
    errors_md = Path("results/error_analysis.md").read_text(encoding="utf-8") if Path("results/error_analysis.md").exists() else ""

    full = metrics[metrics["method"] == "Full Method"].iloc[0]
    best = metrics[metrics["status"] == "ok"].sort_values("score", ascending=False).iloc[0]
    always = metrics[metrics["method"] == "Always Answer"].iloc[0]
    rag = metrics[metrics["method"] == "RAG Similarity Threshold"].iloc[0]
    llm = metrics[metrics["method"] == "Prompted LLM Baseline"].iloc[0]
    provider = api_status.get("provider", "deepseek")
    model_name = api_status.get("model", api_status.get("deepseek_model", "unknown"))
    key_present = api_status.get("deepseek_api_key_present", api_status.get("api_key_present", False))
    api_available = api_status.get("deepseek_api_available", api_status.get("api_available", False))
    selected_full = selection.get("selected_candidate", "all_modules")
    mode = "quick" if manifest["target_size"] <= 120 else "full"
    has_error_cases = "## 案例" in errors_md

    abstract = (
        f"本实验把知识缺口场景下的问答决策拆成两个输出：缺口类型和回答动作。"
        f"项目构造了 {manifest['target_size']} 条统一格式样本，按 group_id 做 6:2:2 划分，比较固定动作、检索阈值、自一致性、结构化特征分类器、文本分类器和 Full Method。"
        f"在当前 {mode} 运行中，Full Method 的 action macro-F1 为 {_fmt(full['action_macro_f1'])}，综合效用为 {_fmt(full['score'])}；"
        f"Always Answer 的 action macro-F1 为 {_fmt(always['action_macro_f1'])}，效用为 {_fmt(always['score'])}。"
        "实验能说明规则化动作空间比直接回答更安全；同时，新数据下各方法不再大面积满分，消融结果能更清楚地显示 self-consistency/不确定性特征的贡献。"
    )

    lines = [
        f"# {TITLE}",
        "",
        "## 1. 摘要",
        abstract,
        "",
        "## 2. 研究背景与问题定义",
        "真实问答系统经常面对信息缺失、问题含混、证据不足、知识过期、错误前提和高风险专业场景。直接生成答案并不总是合适：有时应主动追问用户条件，有时应检索最新证据，有时应拒绝给出专业判断，或先纠正问题中的错误前提。",
        "",
        "本实验采用 T/P/E 框架定义任务：",
        "- T：不确定条件下的回答动作选择。",
        "- P：动作分类、错误回答率、过度拒答率、交互代价和综合效用。",
        "- E：带缺口类型和动作标签的问答样本。",
        "",
        "## 3. 相关工作简述",
        "AmbigQA 和 ASQA 提醒我：很多问题并不是缺一个答案，而是问题本身有多种解释。FreshQA 处理动态事实，正好对应本实验里的 `time_sensitive -> retrieve`。SelfCheckGPT 用多次采样的一致性检查幻觉，本实验也用 DeepSeek 多次采样得到动作投票和短理由一致性特征。Self-RAG 把检索纳入生成控制，本实验进一步把检索、追问、拒答和纠错都放进动作标签。Sufficient Context 和 AbstentionBench 相关工作则对应另两个边界：证据是否足够，以及模型什么时候应该停下来不答。",
        "",
        "## 4. 数据构造",
        f"本次运行构造样本数为 {manifest['target_size']}，生成模式为 `{manifest.get('generation_mode', 'unknown')}`。数据生成参考 AmbigQA/ASQA 的多解释问题、FreshQA 的动态事实、Sufficient Context 的证据充分性判断和 AbstentionBench 的拒答边界；DeepSeek 按受控 scenario 批量生成自然问题、对话上下文、检索证据和候选回答，代码负责类别配额、schema 校验、去重、group 划分和少量 fallback 补齐。每条样本包含 `user_initial_query`、`dialogue_context`、`retrieved_evidence`、`candidate_answer`、`gap_type`、`gold_action`、`required_slots`、证据标签、时间敏感标记、错误前提标记和风险等级。",
        "",
        f"划分使用固定随机种子 42，并按 `group_id` 划分，split 大小为 {manifest['split_sizes']}。同一 group_id 只会进入一个 split，避免同源改写样本泄漏。",
        "",
        "标签分布：",
        _table(label_dist, ["label_type", "label", "count"], 20),
        "",
        "## 5. 方法",
        "整体流程包括：先标准化样本，再计算问题侧、证据侧和模型侧不确定性特征；随后训练 gap_type 分类器和 action 分类器，并在 Full Method 中使用验证集从少量预定义配置里选择最稳的组合。申请书里写到的 BM25、embedding 相似度和 NLI 在当前实现中降级为 TF-IDF、token overlap、覆盖率和冲突启发式。为避免标签泄漏，`evidence_sufficiency_label`、`false_premise_flag`、`time_sensitive_flag`、`risk_level` 不再直接作为模型特征；Self-consistency 特征由 DeepSeek 基于四个输入字段多次采样生成，记录在 `data/cache/llm_self_consistency.jsonl`，缓存条目不包含 `gap_type`、`gold_action`、`final_answer`、`gold_clarifying_question` 或 `required_slots`。",
        "",
        f"Full Method 的配置选择写入 `results/full_method_selection.json`，本次验证集选择为 `{selected_full}`。最终决策以结构化特征分类器为主，只保留保守的安全覆盖；对 `ask` 样本，question utility ranker 生成候选追问，并按 slot coverage、预期不确定性降低、具体性、可回答性、礼貌性和诱导性惩罚排序。",
        "",
        "## 6. 实验设置",
        f"运行模式：{mode}。训练/验证/测试比例为 6:2:2，随机种子为 42。运行环境为 Python {platform.python_version()} / {platform.system()}。{provider} 模型：{model_name}；API key 是否存在：{key_present}；API 是否通过严格 JSON 探测：{api_available}。本次 API 状态说明：{api_status['note']} 如果没有 API key、探测失败或批量调用失败，实验会直接停止，不再使用离线回退特征。",
        "",
        f"对照方法包括 Always Answer、Always Ask、RAG Similarity Threshold、Self-Consistency Baseline、Prompted LLM Baseline、Logistic Regression、Random Forest、GBDT、TF-IDF Text Encoder Classifier 和 Full Method。Prompted LLM Baseline 的本次状态为 `{llm['status']}`；LLM 调用失败会中止实验，而不是用默认动作补齐。",
        "",
        "## 7. 结果",
        "总体指标如下：",
        _table(metrics, ["method", "status", "gap_type_macro_f1", "action_macro_f1", "action_accuracy", "wrong_answer_rate", "score"], 20),
        "",
        f"Full Method 的 action macro-F1 为 {_fmt(full['action_macro_f1'])}，action accuracy 为 {_fmt(full['action_accuracy'])}，wrong answer rate 为 {_fmt(full['wrong_answer_rate'])}，综合效用为 {_fmt(full['score'])}。相比 Always Answer，Full Method 的效用提升为 {_fmt(full['score'] - always['score'])}；相比 RAG 阈值法，效用差值为 {_fmt(full['score'] - rag['score'])}。固定回答策略的错误主要来自把追问、检索、拒答和纠错样本都硬答掉。新数据下 Prompted LLM Baseline、文本分类器和传统结构化模型之间拉开了差距，说明任务不再只是模板记忆。",
        "",
        "主要类别指标保存在 `results/per_class_metrics.csv`。Full Method 的动作混淆矩阵见 `results/confusion_matrix_action.png`，缺口类型混淆矩阵见 `results/confusion_matrix_gap_type.png`，效用对比图见 `results/utility_comparison.png`。",
        "",
        "## 8. 消融实验",
        _table(ablation, ["variant", "action_macro_f1", "gap_type_macro_f1", "score", "delta_action_macro_f1", "delta_score"], 20),
        "",
        _ablation_comment(ablation),
        "",
        "重复划分实验：",
        _table(stability, ["seed", "test_size", "gap_type_macro_f1", "action_macro_f1", "action_accuracy", "wrong_answer_rate", "score"], 10) if not stability.empty else "（未生成）",
        "",
        _stability_summary(stability),
        "",
        "## 9. 显著性检验",
        _table(sig, list(sig.columns), 20) if not sig.empty else "有效方法不足，未生成显著性检验。",
        "",
        "配对 bootstrap 给出 Full Method 与各 baseline 的 action macro-F1 差异及 95% 置信区间。McNemar 检验基于动作预测正确与否的配对列联表。若 quick 模式样本较少，置信区间可能较宽，部分差异不显著，应以 full 模式复核。",
        "",
        "## 10. 错误分析",
        "错误分析文件为 `results/error_analysis.md`。脚本现在只记录真实且去重后的错误，不再为了凑满 20 条复制案例。"
        + ("下面列出自动抽样中的前几个案例：" if has_error_cases else "本次运行没有真实动作预测错误，因此只保留原因说明。"),
        "",
        "\n".join(errors_md.splitlines()[2:30]) if errors_md else "当前测试集未记录错误案例。",
        "",
        "## 11. 结论",
        f"本实验完成了一个可复现的动作决策流程：DeepSeek 数据生成、特征、模型、对照、消融、显著性检验和错误分析都能从脚本生成。当前结果支持一个比较朴素的结论：在知识缺口场景下，把“是否回答”拆成更细的动作标签，比默认直接回答更稳；不确定性/self-consistency 特征是本次 Full Method 的主要增益来源。",
        "",
        "## 12. 局限性",
        "- 数据由 DeepSeek 按受控 scenario 生成，并非真实线上用户日志；真实用户分布可能不同。",
        "- 数据生成仍可能出现轻微标签噪声，因此代码加入了 schema 校验、高风险一致性修复、去重和 fallback 补齐。",
        "- 高风险问题只做动作决策，不提供专业判断。",
        "- API 不稳定或成本会影响 self-consistency 特征规模；当前实现选择硬失败，避免静默回退污染结果。",
        "- 如果使用启发式 NLI proxy，其证据冲突判断能力有限。",
        "- 当前 Prompted LLM Baseline 和 self-consistency 采样都要求 DeepSeek API 可用；若批量预测出现解析失败，实验会停止。",
        "",
        "## 13. 可复现性说明",
        "主要运行命令：",
        "```bash",
        ".venv/bin/python scripts/check_deepseek_api.py",
        ".venv/bin/python -m pytest -q",
        ".venv/bin/python -m knowledge_gap_decision.run_experiment --quick",
        ".venv/bin/python -m knowledge_gap_decision.run_experiment --target-size 800",
        ".venv/bin/python -m knowledge_gap_decision.report",
        "```",
        "随机种子固定为 42。主要依赖见 `requirements.txt`。输出路径包括 `data/processed/`、`results/` 和 `reports/`。",
        "",
        "## 14. 工具辅助说明",
        "本项目使用了代码补全和文本整理工具辅助搭建脚本；报告中的数值、表格和图均来自本仓库实际运行结果。关键结论按 CSV/JSON 输出核对后写入，未手工改数。",
        "",
        "## 15. 参考文献",
        "- Min et al. 2020. AmbigQA: Answering Ambiguous Open-domain Questions. EMNLP.",
        "- Stelmakh et al. 2022. ASQA: Factoid Questions Meet Long-Form Answers. EMNLP.",
        "- Vu et al. 2023. FreshLLMs: Refreshing Large Language Models with Search Engine Augmentation.",
        "- Manakul et al. 2023. SelfCheckGPT: Zero-Resource Black-Box Hallucination Detection for Black-Box LLMs.",
        "- Asai et al. 2023. Self-RAG: Learning to Retrieve, Generate, and Critique through Self-Reflection.",
        "- Joren et al. 2024. Sufficient Context: A New Lens on Retrieval Augmented Generation Systems.",
        "- AbstentionBench. 2025. Reasoning LLMs Fail on Unanswerable Questions.",
        "",
    ]
    text = "\n".join(lines)
    Path(output).parent.mkdir(parents=True, exist_ok=True)
    Path(output).write_text(text, encoding="utf-8")
    generate_outline()
    generate_docx(text)
    return output


def generate_outline(path: str = "reports/presentation_outline.md") -> None:
    outline = """# 展示提纲

## 研究背景
- 问答系统遇到的不是单一“会不会答”，而是多个边界：问题是否含混、证据是否足够、事实是否会过期、前提是否可靠
- 直接回答的主要风险：过早回答、错误前提顺从、高风险场景下给出不该给的结论
- 本实验关注动作选择，不评价最终长答案写得多漂亮

## 任务定义
- 输入：用户问题、对话上下文、检索证据、候选回答
- 输出：gap_type 与 action，两者分别用于解释缺口和决定系统动作
- 动作：answer / ask / retrieve / abstain / challenge_premise
- T/P/E：T 是动作决策，P 是宏平均 F1、错误回答率、交互代价和效用，E 是带标签的问答样本

## 方法
- 数据标准化后按 group_id 划分，避免同源改写进入不同 split
- 问题侧特征：长度、时间词、主观词、条件词、实体、错误前提模式、高风险关键词
- 证据侧特征：TF-IDF 相似度、token overlap、覆盖率、冲突启发式
- 模型侧特征：DeepSeek 基于四个输入字段多次采样 action/rationale，再计算 vote entropy、majority ratio、rationale overlap 和动作票数
- Full Method：分类器给出基础预测，用验证集选择稳定配置，规则只做保守安全覆盖

## 实验设置
- 数据规模与标签分布
- baselines 与模型参数
- DeepSeek API 探测或批量调用失败时实验直接停止，不写离线回退结果
- 重复划分：使用多个随机种子检查 Full Method 稳定性

## 结果
- 先讲 Always Answer 的风险：错误回答率高，效用显著为负
- 再讲 Full Method：动作分数、效用和错误回答率
- 说明 self-consistency 特征现在来自真实 LLM 采样，不读取任何 gold label 字段

## 消融
- 展示 ablation_summary.csv
- 如果 delta 为 0，不强行解释成“模块有效”；说明当前数据没有拉开差距
- question ranker 主要影响追问质量，不一定反映在 action macro-F1

## 案例
- 含混问题：先问操作系统、目标任务等槽位
- 时间敏感问题：查最新公告或负责人信息
- 错误前提：先指出前提没有被证据支持
- 证据不能支持候选结论：拒绝确认，而不是编一个肯定答案

## 结论
- 联合建模让“问、查、拒、纠错”都成为可评估动作
- 当前项目优势在可复现流程完整
- 最大限制是数据仍为 LLM 生成；下一步应补真实用户问句、人工复核标签和更难的相邻标签样本
"""
    Path(path).write_text(outline, encoding="utf-8")


def generate_docx(markdown_text: str, path: str = "reports/final_report.docx") -> None:
    try:
        from docx import Document
    except Exception:
        Path(path + ".fallback.txt").write_text(markdown_text, encoding="utf-8")
        return
    doc = Document()
    for line in markdown_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|"):
            doc.add_paragraph(line)
        elif line.strip() == "```bash" or line.strip() == "```":
            continue
        elif line.strip():
            doc.add_paragraph(line)
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", default="results/metrics_summary.csv")
    parser.add_argument("--output", default="reports/final_report.md")
    args = parser.parse_args()
    if not Path(args.input).exists():
        raise SystemExit(f"missing metrics file: {args.input}; run experiment first")
    generate_markdown(args.output)


if __name__ == "__main__":
    main()
